#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
从 LaTeX 源码生成 Word 版论文（论文/main.docx）。

用法（在项目根目录下执行）：
    python 论文/make_word.py all      # 渲染 TikZ 图 -> pandoc 转换 -> docx 后处理
    python 论文/make_word.py render   # 仅渲染 4 幅 TikZ 示意图为 PNG（需 xelatex）
    python 论文/make_word.py convert  # 仅执行 pandoc 转换与 docx 后处理（需先 render）

依赖：MiKTeX xelatex、pandoc、python-docx、PyMuPDF(fitz)。
说明：Word 版由同一 LaTeX 源码机械转换而来，正文、图表与公式以 main.pdf 为准；
      Word 版用于电子版提交中的 Word 副本，已去除承诺书/编号专用页。
"""

import argparse
import re
import subprocess
import sys
from pathlib import Path

BASE = Path(__file__).resolve().parent

SECTION_FILES = [
    "sections/abstract.tex",
    "sections/s1_problem.tex",
    "sections/s2_analysis.tex",
    "sections/s3_assumptions.tex",
    "sections/s4_symbols.tex",
    "sections/s5_data.tex",
    "sections/s6_problem1.tex",
    "sections/s7_problem2.tex",
    "sections/s8_problem3.tex",
    "sections/s9_conclusion.tex",
    "sections/ai_declaration.tex",
    "refs.tex",
    "sections/appendix.tex",
]

# 与 main.tex 保持一致的配色与 TikZ 样式（Word 转换仅用于渲染示意图）
TIKZ_PREAMBLE = r"""
\definecolor{CUMCMBlue}{HTML}{0072B2}
\definecolor{CUMCMOrange}{HTML}{E69F00}
\definecolor{CUMCMGreen}{HTML}{009E73}
\definecolor{CUMCMRed}{HTML}{D55E00}
\definecolor{CUMCMBg}{HTML}{F7F7F7}
\definecolor{CUMCMBorder}{HTML}{CCCCCC}
\definecolor{CUMCMCharcoal}{HTML}{333333}
\definecolor{CUMCMArrow}{HTML}{4D4D4D}
\definecolor{CUMCMGrey}{HTML}{666666}
\tikzset{
  figbox/.style={draw=CUMCMBorder, fill=white, rounded corners=2pt,
                 align=center, font=\zihao{-5}, text=CUMCMCharcoal,
                 minimum height=0.95cm, inner sep=4pt},
  corebox/.style={figbox, draw=CUMCMBlue, line width=0.7pt},
  altbox/.style={figbox, draw=CUMCMOrange, line width=0.7pt},
  outbox/.style={figbox, draw=CUMCMGreen, line width=0.7pt},
  warnbox/.style={figbox, draw=CUMCMRed, line width=0.7pt},
  stagetitle/.style={font=\zihao{-6}\sffamily, text=CUMCMBlue},
  stagetitlealt/.style={font=\zihao{-6}\sffamily, text=CUMCMOrange},
  stagetitleout/.style={font=\zihao{-6}\sffamily, text=CUMCMGreen},
  figarr/.style={-{Stealth[length=2mm]}, draw=CUMCMArrow, thick},
  figannote/.style={font=\zihao{-6}, text=CUMCMGrey},
}
"""


def load_flattened_tex() -> str:
    """读取 main.tex，内联所有 \\input 的子文件，并剔除纸质版前置页条件块。"""
    main = (BASE / "main.tex").read_text(encoding="utf-8")
    # 注意：\\newif\\ifpaperversion 也含子串 \\ifpaperversion，
    # 必须用行首锚定，避免误吞 \\begin{document}。
    main = re.sub(r"(?m)^\\ifpaperversion[\s\S]*?\\fi$", "", main)

    def _inline(m):
        name = m.group(1)
        p = BASE / (name if name.endswith(".tex") else name + ".tex")
        if not p.exists():
            print(f"[make_word] 警告：找不到 \\input 文件 {name}", file=sys.stderr)
            return ""
        return p.read_text(encoding="utf-8")

    main = re.sub(r"\\input\{([^}]+)\}", _inline, main)
    return main


def parse_aux() -> dict:
    aux = (BASE / "main.aux").read_text(encoding="utf-8", errors="ignore")
    return {k: n for k, n, _ in re.findall(r"\\newlabel\{([^}]+)\}\{\{([^}]*)\}\{(\d+)\}", aux)}


def parse_bibitems() -> dict:
    refs = (BASE / "refs.tex").read_text(encoding="utf-8")
    keys = re.findall(r"\\bibitem\{([^}]+)\}", refs)
    return {k: i + 1 for i, k in enumerate(keys)}


def resolve_refs(text: str, labels: dict, bib: dict) -> str:
    """把 \\eqref/\\ref/\\cite 换成 main.aux/refs.tex 中的真实编号，避免 pandoc 自行编号。"""
    text = re.sub(
        r"\\eqref\{([^}]+)\}",
        lambda m: f"({labels.get(m.group(1), '?')})",
        text,
    )
    text = re.sub(
        r"\\ref\{([^}]+)\}",
        lambda m: labels.get(m.group(1), "?"),
        text,
    )
    text = re.sub(
        r"\\cite\{([^}]+)\}",
        lambda m: "[" + ",".join(str(bib.get(k.strip(), "?")) for k in m.group(1).split(",")) + "]",
        text,
    )
    return text


def extract_tikz_figures() -> list:
    """按正文顺序提取 4 幅 TikZ 示意图（block + caption + label）。"""
    figs = []
    for fname in SECTION_FILES:
        text = (BASE / fname).read_text(encoding="utf-8")
        for m in re.finditer(r"\\begin\{tikzpicture\}[\s\S]*?\\end\{tikzpicture\}", text):
            start = text.rfind("\\begin{figure}", 0, m.start())
            end = text.find("\\end{figure}", m.end())
            env = text[start:end]
            cap_m = re.search(r"\\caption\{([^}]*)\}", env)
            lab_m = re.search(r"\\label\{([^}]+)\}", env)
            figs.append(
                {
                    "file": fname,
                    "block": m.group(0),
                    "caption": cap_m.group(1) if cap_m else "",
                    "label": lab_m.group(1) if lab_m else "",
                }
            )
    return figs


def clean_caption(s: str) -> str:
    """清洗 LaTeX 题注文本：保留常用希腊字母/符号，去掉其余命令，用于 Word 题注。"""
    greek = {
        "alpha": "α", "beta": "β", "gamma": "γ", "delta": "δ",
        "epsilon": "ε", "lambda": "λ", "mu": "μ", "nu": "ν",
        "sigma": "σ", "tau": "τ", "phi": "φ", "theta": "θ",
        "omega": "ω", "pi": "π",
    }
    for cmd, ch in greek.items():
        s = s.replace("\\" + cmd, ch)
    s = s.replace("\\times", "×").replace("\\sim", "~")
    s = re.sub(r"\$([^$]*)\$", r"\1", s)  # 去掉 $，保留公式内容
    s = re.sub(r"\\[a-zA-Z]+\*?", "", s)
    s = s.replace("{", "").replace("}", "").replace("\\", "")
    s = re.sub(r"\s+", " ", s).strip()
    return s


def parse_figure_captions() -> list:
    """按正文顺序解析所有 figure 环境的 (label, 清洗后题注)。"""
    out = []
    for fname in SECTION_FILES:
        text = (BASE / fname).read_text(encoding="utf-8")
        for m in re.finditer(r"\\begin\{figure\}[\s\S]*?\\end\{figure\}", text):
            env = m.group(0)
            # 一个 figure 环境可能含多个 minipage（各带 caption+label），全部收集
            for cap_m in re.finditer(r"\\caption\{([^}]*)\}\s*\\label\{([^}]+)\}", env):
                out.append((cap_m.group(2), clean_caption(cap_m.group(1))))
    return out


def parse_table_captions() -> list:
    """按正文顺序解析所有 table 环境的 (label, 清洗后题注)。"""
    out = []
    for fname in SECTION_FILES:
        text = (BASE / fname).read_text(encoding="utf-8")
        for m in re.finditer(r"\\begin\{table\}[\s\S]*?\\end\{table\}", text):
            env = m.group(0)
            for cap_m in re.finditer(r"\\caption\{([^}]*)\}\s*\\label\{([^}]+)\}", env):
                out.append((cap_m.group(2), clean_caption(cap_m.group(1))))
    return out


def render_tikz(figs: list) -> dict:
    """把每幅 TikZ 示意图编译为独立 PDF 并栅格化为 PNG，返回 {label: png_path}。"""
    build_dir = BASE / ".word_build"
    build_dir.mkdir(exist_ok=True)
    rendered = {}
    try:
        import fitz  # PyMuPDF
    except ImportError:
        print("[make_word] 需要 PyMuPDF（pip install pymupdf）以栅格化示意图", file=sys.stderr)
        raise

    for i, fig in enumerate(figs):
        name = f"tikz_{i:02d}"
        tex_path = build_dir / f"{name}.tex"
        tex = (
            "\\documentclass[border=6pt]{standalone}\n"
            "\\usepackage{ctex}\n"
            "\\usepackage{amsmath,amssymb}\n"
            "\\usepackage{tikz}\n"
            "\\usetikzlibrary{arrows.meta,positioning,calc,shapes.geometric}\n"
            + TIKZ_PREAMBLE
            + "\\begin{document}\n"
            + fig["block"]
            + "\n\\end{document}\n"
        )
        tex_path.write_text(tex, encoding="utf-8")
        r = subprocess.run(
            ["xelatex", "-interaction=nonstopmode", "-halt-on-error", tex_path.name],
            cwd=str(build_dir),
            capture_output=True,
            text=True,
            encoding="utf-8",
        )
        if r.returncode != 0:
            print(f"[make_word] xelatex 编译失败：{name}（label={fig['label']}）", file=sys.stderr)
            print((r.stdout or "")[-2000:], file=sys.stderr)
            raise RuntimeError(f"TikZ 渲染失败: {fig['label']}")
        pdf_path = build_dir / f"{name}.pdf"
        png_path = build_dir / f"{name}.png"
        doc = fitz.open(str(pdf_path))
        pix = doc[0].get_pixmap(matrix=fitz.Matrix(3.0, 3.0), alpha=False)
        pix.save(str(png_path))
        rendered[fig["label"]] = png_path
        print(f"[make_word] 渲染完成：{fig['label']} -> {png_path.name}")
    return rendered


def convert_to_docx(rendered_tikz: dict):
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.shared import Cm, Mm, Pt, RGBColor

    labels = parse_aux()
    bib = parse_bibitems()
    flat = load_flattened_tex()
    flat = resolve_refs(flat, labels, bib)
    # 参考文献标题（refs.tex 的 thebibliography 不产生标题，pandoc 需要显式章节）
    flat = flat.replace(
        "\\begin{thebibliography}{99}",
        "\\section*{参考文献}\n\\begin{thebibliography}{99}",
    )

    eq_count = len(re.findall(r"\\begin\{equation\}", flat))

    flat_path = BASE / "main_word.tex"
    flat_path.write_text(flat, encoding="utf-8")

    raw_docx = BASE / "main_word_raw.docx"
    r = subprocess.run(
        ["pandoc", "main_word.tex", "-o", "main_word_raw.docx", "--resource-path=."],
        cwd=str(BASE),
        capture_output=True,
        text=True,
        encoding="utf-8",
    )
    if r.returncode != 0:
        print(r.stdout, file=sys.stderr)
        print(r.stderr, file=sys.stderr)
        raise RuntimeError("pandoc 转换失败")

    doc = Document(str(raw_docx))

    # ---------- 页面与字体 ----------
    for sec in doc.sections:
        sec.page_width, sec.page_height = Mm(210), Mm(297)
        sec.left_margin = sec.right_margin = Cm(2.5)
        sec.top_margin = sec.bottom_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")
    for st in doc.styles:
        if st.type != WD_STYLE_TYPE.PARAGRAPH or not re.fullmatch(r"Heading [1-3]", st.name or ""):
            continue
        st.font.name = "Times New Roman"
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "黑体")

    def _set_caption(p, prefix, text):
        for r in list(p.runs):
            r._r.getparent().remove(r._r)
        r1 = p.add_run(f"{prefix} ")
        r1.bold = True
        p.add_run(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---------- 图题注：按位置配对图像与题注并重写编号/文本 ----------
    tikz_labels = list(rendered_tikz.keys())
    fig_caps = parse_figure_captions()
    photo_caps = [(lab, labels[lab], cap) for lab, cap in fig_caps if lab not in tikz_labels]
    paras = list(doc.paragraphs)
    img_idx = [k for k, p in enumerate(paras) if p.style.name == "Captioned Figure"]
    cap_idx = [k for k, p in enumerate(paras) if p.style.name == "Image Caption"]
    consumed = set()
    for i, k in enumerate(img_idx):
        if i >= len(photo_caps):
            break
        lab, num, cap = photo_caps[i]
        cand = [j for j in cap_idx if j > k and j not in consumed]
        if cand:
            ci = cand[0]
            consumed.add(ci)
            _set_caption(paras[ci], f"图{num}", cap)
        else:
            # 缺题注时补插一个
            new_p = doc.add_paragraph()
            new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_caption(new_p, f"图{num}", cap)
            paras[k]._p.addnext(new_p._p)
            print(f"[make_word] 补插图{num} 题注（{cap}）")
    for j in cap_idx:
        if j not in consumed:
            paras[j]._p.getparent().remove(paras[j]._p)

    # ---------- 表题注编号 ----------
    tab_order = [(k, labels[k]) for k in labels if k.startswith("tab:")]
    tab_caps = parse_table_captions()
    tab_text = {lab: cap for lab, cap in tab_caps}
    tc_paras = [p for p in doc.paragraphs if p.style.name == "Table Caption"]
    for i, p in enumerate(tc_paras):
        if i < len(tab_order):
            lab, num = tab_order[i]
            cap = tab_text.get(lab, p.text)
            _set_caption(p, f"表{num}", cap)
        else:
            print(f"[make_word] 警告：表题注多于表标签（{len(tc_paras)} vs {len(tab_order)}）", file=sys.stderr)

    # ---------- 补插 TikZ 示意图（置于其首次引用段之后） ----------
    tikz_order = [(k, labels[k]) for k in tikz_labels if k in labels]
    for label, num in tikz_order:
        caption = next((f["caption"] for f in extract_tikz_figures() if f["label"] == label), "")
        target = None
        for p in doc.paragraphs:
            if f"图 {num}" in p.text or f"图{num}" in p.text:
                target = p
                break
        if target is None:
            print(f"[make_word] 警告：未找到图{num} 的引用位置，跳过插入", file=sys.stderr)
            continue
        img_par = doc.add_paragraph()
        img_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_par.add_run().add_picture(str(rendered_tikz[label]), width=Cm(15.0))
        cap_par = doc.add_paragraph()
        cap_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_par.add_run(f"图{num} {caption}")
        target._p.addnext(cap_par._p)
        target._p.addnext(img_par._p)
        print(f"[make_word] 已插入图{num}（{caption}）")

    # ---------- 公式编号（按正文顺序 1..N） ----------
    eq_paras = [p for p in doc.paragraphs if p._p.findall(qn("m:oMathPara"))]
    if eq_count == len(eq_paras):
        for i, p in enumerate(eq_paras, 1):
            p.paragraph_format.tab_stops.add_tab_stop(Cm(16.0), WD_TAB_ALIGNMENT.RIGHT)
            p.add_run(f"\t({i})")
    else:
        print(
            f"[make_word] 警告：公式数不匹配（源码 {eq_count} vs docx {len(eq_paras)}），跳过公式编号",
            file=sys.stderr,
        )

    # ---------- 摘要页标题与"摘  要"标题 ----------
    title_text = "基于YOLOv8与极值理论的集装箱破损智能检测模型"
    title_p = next((p for p in doc.paragraphs if p.text.strip() == title_text), None)
    if title_p is not None:
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in title_p.runs:
            r.bold = True
            r.font.size = Pt(16)
            r._r.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "黑体")
        abs_p = doc.add_paragraph()
        abs_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = abs_p.add_run("摘  要")
        run.bold = True
        run.font.size = Pt(14)
        run._r.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "黑体")
        title_p._p.addnext(abs_p._p)

    # ---------- 关键词标签 ----------
    kw_p = next(
        (p for p in doc.paragraphs if "集装箱破损检测" in p.text and p.text.strip().startswith("集装箱")),
        None,
    )
    if kw_p is not None and not kw_p.text.strip().startswith("关键词"):
        r_kw = kw_p.add_run("关键词：")
        r_kw.bold = True
        kw_p._p.insert(0, r_kw._r)

    # ---------- 附录标题 ----------
    appendix_names = {
        "主要源程序": "附录A 主要源程序",
        "极值理论判别的推导补充": "附录B 极值理论判别的推导补充",
        "WD-Focal Loss 推导补充": "附录C WD-Focal Loss 推导补充",
    }
    for p in doc.paragraphs:
        name = p.text.strip()
        if p.style.name == "Heading 1" and name in appendix_names:
            for r in list(p.runs):
                r._r.getparent().remove(r._r)
            p.add_run(appendix_names[name])

    out = BASE / "main.docx"
    doc.save(str(out))
    print(f"[make_word] Word 版已生成：{out}")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("stage", choices=["all", "render", "convert"], default="all", nargs="?")
    args = ap.parse_args()

    figs = extract_tikz_figures()
    print(f"[make_word] 识别到 {len(figs)} 幅 TikZ 示意图")
    rendered = {}
    if args.stage in ("all", "render"):
        rendered = render_tikz(figs)
    else:
        for i, fig in enumerate(figs):
            png = BASE / ".word_build" / f"tikz_{i:02d}.png"
            if png.exists():
                rendered[fig["label"]] = png
    if not rendered:
        print("[make_word] 未找到已渲染的示意图，请先执行 render 阶段", file=sys.stderr)
        sys.exit(1)
    if args.stage in ("all", "convert"):
        convert_to_docx(rendered)


if __name__ == "__main__":
    main()
